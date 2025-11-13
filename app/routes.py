import io
from flask import Flask, render_template, request, send_file, url_for, redirect, session, jsonify # type: ignore 
from functools import wraps
import mysql.connector # type: ignore
from datetime import datetime, timedelta
import json
from app import app
import base64
from docx import Document
import io
from app.conexao import config, key


@app.route("/")
def index():
    return redirect("/home")


#Etapas de agendamento
@app.route('/agendamento')
def agendar():
    return render_template('agendamento.html')

@app.route('/home')
def home():
    google_maps_api_key = key
    return render_template('home.html', key=google_maps_api_key)

@app.route('/consulta')
def consulta():
    if 'user' not in session:
        return redirect('/login')
    return render_template('consulta.html')

@app.route('/stands')
def stands():
    google_maps_api_key = key
    return render_template("stands.html", key=google_maps_api_key)

@app.route('/cenografia')
def cenografia():
    google_maps_api_key = key
    return render_template('cenografia.html', key=google_maps_api_key)

@app.route('/quiosques')
def quiosques():
    google_maps_api_key = key
    return render_template('quiosques.html', key=google_maps_api_key)

@app.route('/login', methods=['GET', 'POST'])
def login():
    connection = mysql.connector.connect(**config)
    cursor = connection.cursor()
    if request.method == 'POST':
        user = request.form['user']
        pwd = request.form['pwd']
        cursor.execute("SELECT * FROM user WHERE user=%s AND pwd=%s", (user, pwd))
        result = cursor.fetchone()
        connection.close()
        if result:
            session['user'] = user
            return redirect('/links')
        else:
            return "Usuário ou senha inválidos"
    return render_template('login.html')

@app.route('/links')
def links():
    if 'user' not in session:
        return redirect('/login')
    return render_template('links.html')  # Sua tela de links protegida

@app.route("/politica-privacidade")
def politica_privacidade():
    return render_template("politica-privacidade.html")


# agendamento de usuário
@app.route('/agendamento', methods=['POST'])
def agendamento():
         if request.method == 'POST':
            nome = request.form.get('nome')
            telefone = request.form.get('telefone')
            email_cliente = request.form.get('email_cliente')
            data = request.form.get('data')
            aceite_termo = request.form.get('aceite_termo')

            #CAMPOS BRIEFING STEP1
            endereco = request.form.get('Endereco')
            razao_social = request.form.get('Razao_social')
            nome_fantasia = request.form.get('Nome_Fantasia')
            site = request.form.get('Site')
            data_entrega = request.form.get('data_entrega')
            nome_responsavel = request.form.get('Nome_Responsavel')
            telefone_responsavel = request.form.get('telefone_responsavel')
            email_empresa = request.form.get('Email')            
            evento = request.form.get('Evento')
            local = request.form.get('Local')
            stand =  request.form.get('Stand')
            data_evento = request.form.get('data_evento')
            informacoes_adicionas = request.form.get('Informacoes_adicionas')
            #STEP2
            valor_verba = request.form.get('Valor_verba')
            contato = request.form.get('Contato')
            data_atual = request.form.get('data_atual')
            espaco_stand = request.form.get('espaco_stand')
            medida_frente = request.form.get('medida_Frente')
            medida_fundo = request.form.get('medida_Fundo')
            area_total = request.form.get('Area_total')
            estilo_construcao = request.form.get('estilo_construcao')
            adicionais = request.form.get('Adicionais')
            cor_mdf = request.form.get('cor_mdf_input')
            cor_forracao = request.form.get('cor_forracao')
            sala = request.form.get('sala')
            area_exposicao = request.form.get('area_exposicao')
            cores_empresa = request.form.get('Cores_empresa')
            produtos = request.form.get('produtos')
            listamobiliario = request.form.getlist('listaMobiliario')
            # Transformar lista de mobiliário em string com quebras de linha
            listamobiliario_str = "\n".join(listamobiliario) if listamobiliario else ""

            campos = {
                "Endereço": endereco,
                "Razão Social": razao_social,
                "Nome Fantasia": nome_fantasia,
                "Data de Entrega": data_entrega,
                "Nome do Responsável": nome_responsavel,
                "Telefone do Responsável": telefone_responsavel,
                "Email": email_empresa,
                "Email do cliente": email_cliente,
                "Site": site,
                "Evento": evento,
                "Local": local,
                "Stand": stand,
                "Valor da verba": valor_verba,
                "Data do Evento": data_evento,
                "Informações adicionais": informacoes_adicionas,
                "Contato": contato,
                "Data atual": data_atual,
                "Espaço stand": espaco_stand,
                "Estilo construção": estilo_construcao, 
                "Medida Frente": medida_frente,
                "Medida Fundo": medida_fundo,
                "Área Total": area_total,
                "Adicionais": adicionais,
                "Cor MDF": cor_mdf,
                "Cor Forração": cor_forracao,
                "Sala": sala,
                "Área de exposição": area_exposicao,
                "Cores da Empresa": cores_empresa,
                "Produtos": produtos,
                "Lista de Mobiliário": listamobiliario_str
                }


                # Criar documento
            document = Document()
            document.add_heading("Briefing de Agendamento"+"\n", level=1)

            # Adiciona todos os campos no documento
            for chave, valor in campos.items():
                if valor:
                    document.add_paragraph(f"{chave}: {valor}")

            # Depois do loop, cria o arquivo DOCX em memória
            doc_io = io.BytesIO()
            document.save(doc_io)
            doc_io.seek(0)
            descricao_projeto = doc_io.read()  # aqui sim o documento tem conteúdo

            # Lê o arquivo uploadado
            doc_briefing = request.files.get('doc_briefing')
            doc_briefing_data = doc_briefing.read() if doc_briefing else None





            # Conexão com o banco
            connection = mysql.connector.connect(**config)
            cursor = connection.cursor()

            try:  
                # Inserção no banco (barbeiro)
                sql = "INSERT INTO agendamento (nome, telefone,email , data, descricao_projeto, aceite_termo , doc_briefing) VALUES ( %s, %s, %s, %s, %s, %s, %s)"
                valores = (nome, telefone,email_cliente , data, descricao_projeto, aceite_termo, doc_briefing_data)
                cursor.execute(sql, valores)
                connection.commit()
                
                print("Agendamento feito com sucesso!")
                return render_template('agendamento.html', mensagem_sucesso="Agendamento feito com sucesso!")

            
            except Exception as e:
                connection.rollback()
                return f"Erro: {e}"
            finally:
                cursor.close()
                connection.close()

# Rota para buscar agendamentos do banco de dados
@app.route('/api/agendamentos', methods=['GET'])
def api_agendamentos():
    connection = mysql.connector.connect(**config)
    cursor = connection.cursor()

    query = """
    SELECT id, nome, telefone, email, DATE_FORMAT(data, '%d/%m/%Y') as data,
           descricao_projeto, doc_briefing,
           CASE
               WHEN aceite_termo = 'on' THEN 'sim'
               ELSE 'não'
           END AS aceite_termo
    FROM agendamento;"""

    cursor.execute(query)  # sem f-string aqui
    agendamentos = cursor.fetchall()
    connection.close()

    return jsonify([
        {
            "id": row[0],
            "nome": row[1],
            "telefone": row[2],
            "email": row[3],
            "data": row[4],
            "descricao_projeto": base64.b64encode(row[5]).decode('utf-8') if row[5] else None,
            "doc_briefing": base64.b64encode(row[6]).decode('utf-8') if row[6] else None,
            "aceite_termo": row[7]
        }
        for row in agendamentos
    ])

@app.route("/download_projeto/<int:id>", methods=["GET"])
def download_projeto(id):
    connection = mysql.connector.connect(**config)
    cursor = connection.cursor()
    cursor.execute("SELECT descricao_projeto FROM agendamento WHERE id = %s", (id,))
    resultado = cursor.fetchone()
    connection.close()

    if resultado and resultado[0]:
        arquivo_binario = resultado[0]
        arquivo_io = io.BytesIO(arquivo_binario)

        return send_file(
            arquivo_io,
            as_attachment=True,
            download_name=f"projeto_{id}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        return "Documento do projeto não encontrado", 404



@app.route("/download/<int:id>", methods=["GET"])
def download_arquivo(id):
    connection = mysql.connector.connect(**config)
    cursor = connection.cursor()
    
    cursor.execute("SELECT doc_briefing FROM agendamento WHERE id = %s", (id,))
    resultado = cursor.fetchone()
    connection.close()

    if resultado and resultado[0]:  
        arquivo_binario = resultado[0]

        # Criar um objeto de Bytes
        arquivo_io = io.BytesIO(arquivo_binario)

        return send_file(
            arquivo_io, 
            as_attachment=True, 
            download_name=f"briefing_{id}.docx",  # Nome do arquivo baixado
            mimetype="application/docx"  # Ajuste conforme necessário
        )
    else:
        return "Arquivo não encontrado", 404

@app.route('/items/<int:id>', methods=['PUT'])
def update_item(id):
    try:
        connection = mysql.connector.connect(**config)
        cursor = connection.cursor()

        # Verifica se o registro existe
        cursor.execute("SELECT id FROM agendamento WHERE id = %s", (id,))
        resultado = cursor.fetchone()
        if not resultado:
            return jsonify({"message": "registro não encontrado"}), 404

        # Recebe os dados do JSON enviado pelo JS
        data = request.get_json()
        nome = data.get("nome")
        email = data.get("email")
        telefone = data.get("telefone")
        data_evento = data.get("data")  # opcional

        # Atualiza o registro
        cursor.execute("""
            UPDATE agendamento
            SET nome = %s, email = %s, telefone = %s, data = %s
            WHERE id = %s
        """, (nome, email, telefone, data_evento, id))

        connection.commit()

        return jsonify({"message": "registro atualizado com sucesso"})

    except mysql.connector.Error as err:
        return jsonify({"message": f"Erro: {err}"}), 500

    finally:
        cursor.close()
        connection.close()




@app.route('/items/<int:id>', methods=['DELETE'])
def delete_item(id):
    try:
        connection = mysql.connector.connect(**config)
        cursor = connection.cursor()

        # Verifica se o registro existe
        cursor.execute("SELECT id FROM agendamento WHERE id = %s", (id,))
        resultado = cursor.fetchone()
        if not resultado:
            return jsonify({"message": "registro não encontrado"}), 404

        # Deleta o registro
        cursor.execute("DELETE FROM agendamento WHERE id = %s", (id,))
        connection.commit()

        return jsonify({"message": "registro deletado com sucesso"})

    except mysql.connector.Error as err:
        return jsonify({"message": f"Erro: {err}"}), 500

    finally:
        cursor.close()
        connection.close()



